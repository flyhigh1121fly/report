"""动态文档章节构建器。

从封面模板出发，根据前端 image_config 程序化构建每个内容章节，
追加到封面之后、分析与讨论之前。

章节结构（最终顺序）:
    章节标题（加粗、宋体 14pt）
    → 引导文字（宋体 10.5pt）
    → 图片（单图居中段落 / 多图并排无边框表格）
    → 子图标注（表格内或图片下方）
    → 图题（居中）
    → VLM 分析文本（宋体 10.5pt）
"""

from __future__ import annotations

import re
from pathlib import Path
from typing import Any, List

from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ─────────────────────
# 常量
# ─────────────────────
FONT_NAME_CN = "宋体"
FONT_SIZE_BODY_PT = 10.5  # 五号
FONT_SIZE_TITLE_PT = 14.0
WORD_IMAGE_WIDTH_CM = 6.62  # 单张图片全宽
SIDE_BY_SIDE_WIDTH_CM = 7.0  # 并排时单张宽度
IMAGES_PER_ROW = 2


# ─────────────────────
# Anchor 代理：统一段落和表格的 XML 操作
# ─────────────────────
class _AnchorProxy:
    """包装 XML 元素，使调用方可以通过 ._p 访问。

    无论是段落（w:p）还是表格（w:tbl），都用 ._p 暴露 XML 元素，
    调用方 `anchor._p.addnext(new_element)` 始终能工作。
    """
    def __init__(self, xml_element):
        self._p = xml_element


# ─────────────────────
# 工具函数
# ─────────────────────
def set_run_font(run, font_name: str, size_pt: float) -> None:
    """设置 run 的中文字体和大小。"""
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    r = run._element.rPr
    if r is None:
        r = run._element.get_or_add_rPr()
    rFonts = r.rFonts
    if rFonts is None:
        rFonts = r.get_or_add_rFonts()
    rFonts.set(qn("w:eastAsia"), font_name)
    rFonts.set(qn("w:ascii"), font_name)
    rFonts.set(qn("w:hAnsi"), font_name)


def _insert_paragraph_after(doc: Document, anchor, text: str = "",
                            bold: bool = False, center: bool = False,
                            font_size: float = FONT_SIZE_BODY_PT) -> Any:
    """在 anchor 之后插入新段落，返回新段落对象。"""
    new_p = doc.add_paragraph()
    anchor._p.addnext(new_p._p)
    run = new_p.add_run(text)
    run.bold = bold
    set_run_font(run, FONT_NAME_CN, font_size)
    if center:
        new_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return new_p


def _remove_table_borders(table) -> None:
    """移除表格所有边框。"""
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = tbl.get_or_add_tblPr()
    borders = OxmlElement("w:tblBorders")
    for name in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{name}")
        el.set(qn("w:val"), "none")
        el.set(qn("w:sz"), "0")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "auto")
        borders.append(el)
    existing = tblPr.find(qn("w:tblBorders"))
    if existing is not None:
        tblPr.remove(existing)
    tblPr.append(borders)


def _set_cell_vertical_alignment(cell, align: str = "center") -> None:
    """设置单元格垂直对齐。"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    vAlign = OxmlElement("w:vAlign")
    vAlign.set(qn("w:val"), align)
    tcPr.append(vAlign)


def _set_cell_shading(cell, color: str) -> None:
    """设置单元格背景色。"""
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def _format_caption_label(label: str, mark: str) -> str:
    """格式化子图标注文本，避免标号重复。

    前端传入的 label 可能已是 "(a) 显微组织 100x" 格式（含英文括号标记）。
    此函数将英文括号统一为中文括号，避免与额外追加的标记重复。
    """
    # 如果 label 以 (x) 或 （x） 开头，替换为中文括号版本
    cleaned = re.sub(r'^\(([a-z])\)\s*', r'（\1）', label)
    if cleaned.startswith("（"):
        return cleaned
    # label 无标记前缀，手动添加
    return f"（{mark}）{label}"


# ─────────────────────
# 构建器
# ─────────────────────
def build_section_title(doc: Document, anchor: Any, title: str) -> Any:
    """构建章节标题段落（加粗、宋体 14pt），插入到 anchor 之后。"""
    p = _insert_paragraph_after(doc, anchor, title,
                                bold=True, center=False,
                                font_size=FONT_SIZE_TITLE_PT)
    return p


def build_intro_text(doc: Document, anchor: Any, text: str) -> Any:
    """构建引导文字段落。如果 text 为空或 "none"，返回原 anchor。"""
    if not text or text == "none":
        return anchor
    return _insert_paragraph_after(doc, anchor, text,
                                   bold=False, center=False)


def _create_image_table(doc: Document, image_slots: List[dict],
                        width_cm: float) -> Any:
    """创建并排图片表格。

    每行 2 张图片，图片下方各有一行标注。返回 table 对象。
    """
    n = len(image_slots)
    row_pairs = (n + IMAGES_PER_ROW - 1) // IMAGES_PER_ROW

    table = doc.add_table(rows=row_pairs * 2, cols=IMAGES_PER_ROW)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _remove_table_borders(table)

    for slot_idx, slot in enumerate(image_slots):
        pair_idx = slot_idx // IMAGES_PER_ROW
        col = slot_idx % IMAGES_PER_ROW
        mark = chr(ord('a') + slot_idx)

        # 图片行
        img_cell = table.cell(pair_idx * 2, col)
        img_cell.text = ""
        _set_cell_vertical_alignment(img_cell, "center")
        p_img = img_cell.paragraphs[0]
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_img = p_img.add_run()
        run_img.add_picture(str(slot["path"]), width=Cm(width_cm))

        # 标注行
        cap_cell = table.cell(pair_idx * 2 + 1, col)
        cap_cell.text = ""
        _set_cell_vertical_alignment(cap_cell, "center")
        p_cap = cap_cell.paragraphs[0]
        p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap_text = _format_caption_label(slot['label'], mark)
        run_cap = p_cap.add_run(cap_text)
        set_run_font(run_cap, FONT_NAME_CN, FONT_SIZE_BODY_PT)

    # 清空多余的单元格（奇数张图时）
    if n % IMAGES_PER_ROW != 0:
        last_pair = row_pairs - 1
        empty_start = n % IMAGES_PER_ROW
        for c in range(empty_start, IMAGES_PER_ROW):
            for r in range(2):
                cell = table.cell(last_pair * 2 + r, c)
                cell.text = ""

    return table


def build_data_table(doc: Document, anchor: Any,
                     headers: List[str], rows: List[List[str]]) -> Any:
    """构建数据表格（带边框、表头加粗、浅灰底色），插入到 anchor 之后。"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 表头行
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        set_run_font(run, FONT_NAME_CN, FONT_SIZE_BODY_PT)
        _set_cell_shading(cell, "D9E2F3")

    # 数据行
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.cell(r_idx + 1, c_idx)
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(val))
            set_run_font(run, FONT_NAME_CN, FONT_SIZE_BODY_PT)

    anchor._p.addnext(table._tbl)
    return table


def build_complete_section(doc: Document, anchor: Any,
                           section_title: str,
                           image_slots: List[dict],
                           analysis_text: str,
                           figure_caption: str = "",
                           intro_text: str = "",
                           image_width_cm: float = WORD_IMAGE_WIDTH_CM) -> Any:
    """构建一个完整的章节。

    最终文档顺序：标题 → 引导文字 → 图片 → 图题 → 分析文本

    使用正序插入 + XML 元素追踪，确保 anchor 链始终指向最后元素。

    Returns
    -------
    _AnchorProxy — 包装了最后一个 XML 元素，有 ._p 属性可供 addnext
    """
    if not image_slots:
        return anchor

    # ── 1. 章节标题 ──
    title_p = build_section_title(doc, anchor, section_title)
    last_xml = title_p._p

    # ── 2. 引导文字 ──
    if intro_text and intro_text != "none":
        intro_p = _insert_paragraph_after(doc, _AnchorProxy(last_xml), intro_text)
        last_xml = intro_p._p

    # ── 3. 表格 + 图片 ──
    table_slots = [s for s in image_slots if s.get("table_data")]
    image_slots_only = [s for s in image_slots if not s.get("table_data")]

    # 表格标题（在表格上方）
    if figure_caption and table_slots:
        caption_p = doc.add_paragraph()
        last_xml.addnext(caption_p._p)
        caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_cap = caption_p.add_run(figure_caption)
        set_run_font(run_cap, FONT_NAME_CN, FONT_SIZE_BODY_PT)
        last_xml = caption_p._p

    # 渲染提取到的数据表格
    for slot in table_slots:
        td = slot["table_data"]
        tbl = build_data_table(doc, _AnchorProxy(last_xml),
                               headers=td["headers"], rows=td["rows"])
        last_xml = tbl._tbl

    # 渲染图片
    if len(image_slots_only) == 1:
        # 单张图片：不用表格，直接居中段落，不加子图标注
        slot = image_slots_only[0]

        # 图片段落
        img_p = doc.add_paragraph()
        last_xml.addnext(img_p._p)
        img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_img = img_p.add_run()
        run_img.add_picture(str(slot["path"]), width=Cm(image_width_cm))
        last_xml = img_p._p

        # 标注段落（单张不加 a/b 前缀）
        cap_p = doc.add_paragraph()
        last_xml.addnext(cap_p._p)
        cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_cap = cap_p.add_run(slot["label"])
        set_run_font(run_cap, FONT_NAME_CN, FONT_SIZE_BODY_PT)
        last_xml = cap_p._p
    elif len(image_slots_only) > 1:
        # 多张图片：并排表格
        per_image_width = SIDE_BY_SIDE_WIDTH_CM
        table = _create_image_table(doc, image_slots_only, per_image_width)
        last_xml.addnext(table._tbl)
        last_xml = table._tbl

    # 图片标题（在图片下方，仅当有图片且无表格时显示）
    if figure_caption and not table_slots:
        caption_p = doc.add_paragraph()
        last_xml.addnext(caption_p._p)
        caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_cap = caption_p.add_run(figure_caption)
        set_run_font(run_cap, FONT_NAME_CN, FONT_SIZE_BODY_PT)
        last_xml = caption_p._p

    # ── 5. 分析文本 ──
    if analysis_text and analysis_text != "none":
        analysis_p = doc.add_paragraph()
        last_xml.addnext(analysis_p._p)
        run_ana = analysis_p.add_run(analysis_text)
        set_run_font(run_ana, FONT_NAME_CN, FONT_SIZE_BODY_PT)
        last_xml = analysis_p._p

    return _AnchorProxy(last_xml)
