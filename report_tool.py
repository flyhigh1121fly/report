from __future__ import annotations

import base64
import io
import json
import os
import queue
import re
import shutil
import tempfile
import threading
import traceback
import uuid
from datetime import datetime
from pathlib import Path
from typing import Dict, Tuple, Any, List

import requests
from flask import Flask, request as flask_request, Response, send_file, send_from_directory
from PIL import Image, ImageOps
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from dotenv import load_dotenv

from section_builders import build_complete_section, build_section_title

load_dotenv()


# =========================
# API 配置（从 .env 读取）
# =========================
BASE_URL = os.environ.get("BASE_URL", "http://122.224.109.54:20001/spiritx-api/v1")
API_KEY = os.environ.get("API_KEY", "")

# 动态组装开关：true 使用封面模板 + 程序化构建章节，false 使用旧模板
USE_DYNAMIC_ASSEMBLY = os.environ.get("USE_DYNAMIC_ASSEMBLY", "true").lower() == "true"

VLM_MODEL = "Qwen/Qwen3-VL-32B-Instruct"
LLM_MODEL = "Qwen/Qwen3-32B"

# =========================
# 知识库配置（从 .env 读取）
# =========================
KB_AUTH_URL = os.environ.get("KB_AUTH_URL", "")
KB_API_URL = os.environ.get("KB_API_URL", "")
KB_USERNAME = os.environ.get("KB_USERNAME", "")
KB_PASSWORD = os.environ.get("KB_PASSWORD", "")

# =========================
# 输出目录
# =========================
JOB_ROOT = Path.home() / "Desktop" / "gradio_report_jobs"

# =========================
# 图片规则
# =========================
TARGET_W_PX = 782
TARGET_H_PX = 591
TARGET_SIZE = (TARGET_W_PX, TARGET_H_PX)

WORD_IMAGE_WIDTH_CM = 6.62  # 插入Word宽度（cm）

# =========================
# 字体要求：替换进去的文本统一宋体五号
# =========================
FONT_NAME_CN = "宋体"
FONT_SIZE_5_PT = 10.5  # 五号

# =========================
# 文本字段（用户填写）
# =========================
TEXT_FIELDS = [
    ("项目名称", "{{项目名称}}"),
    ("背景概述", "{{背景概述}}"),
]
OVERVIEW_SYNC_PLACEHOLDERS = ["{{背景概述}}", "{{概述}}"]

# =========================
# 图片槽位尺寸规则：1-4 章节固定尺寸（782×591），其他原样
# =========================
FIXED_SIZE_SECTIONS = {"1", "2", "3", "4"}

# =========================
# 模型生成替换目标占位符
# =========================
PH_BASIC = "{{基本图片分析}}"
PH_MACRO = "{{宏观断口图片分析}}"
PH_MICRO = "{{微观断口图片分析}}"
PH_MICROSTRUCT = "{{显微组织图片分析}}"
PH_CHEM = "{{化学分析}}"
PH_PERF = "{{性能验证}}"
PH_DISCUSSION = "{{分析与讨论}}"
PH_CONCLUSION = "{{结论}}"
PH_ABSTRACT = "{{报告摘要}}"
PH_KEYWORDS = "{{关键词}}"

# =========================
# 章节类型检测 + 动态提示词模板
# =========================
# 按关键词从具体到通用的顺序匹配
_TYPE_KEYWORDS: list[tuple[str, list[str]]] = [
    ("metallography", ["金相", "显微组织", "组织分析"]),
    ("chemical",      ["化学", "成分", "EDS", "eds", "光谱", "能谱"]),
    ("mechanical",    ["硬度", "力学", "拉伸", "性能测试", "性能验证", "性能"]),
    ("sem",           ["扫描电镜", "SEM", "sem", "微观断口"]),
    ("macro",         ["宏观", "体视", "目视"]),
    ("overview",      ["概述", "外观", "断裂件", "断裂位置"]),
]

TABLE_TYPES = {"chemical", "mechanical"}  # 这两类用"表X"，其余用"图X"


def detect_section_type(title: str, *, is_first: bool = False) -> str:
    """根据章节标题关键词检测分析类型。首个章节默认为 overview。"""
    if is_first:
        return "overview"
    clean = re.sub(r"^[\d.]+\s*", "", title)
    for sec_type, keywords in _TYPE_KEYWORDS:
        if any(kw in clean for kw in keywords):
            return sec_type
    return "generic"


# VLM 观察任务提示词（按分析类型选择）
VLM_TYPE_PROMPTS = {
    "overview": (
        "观察{section_title}，重点描述{slot_labels}，"
        "识别可见的断裂位置、结构特征与异常点。"
    ),
    "macro": (
        "观察{section_title}中的{slot_labels}，"
        "描述断口粗糙度、边缘特征、可能的裂纹起始迹象、"
        "以及断裂面的宏观形貌特征。"
    ),
    "sem": (
        "观察{section_title}中的{slot_labels}，"
        "描述裂纹源区、疲劳条带、扩展区与瞬断区的形貌差异，"
        "以及微观断口的典型特征。"
    ),
    "metallography": (
        "观察{section_title}下的{slot_labels}，"
        "描述组织均匀性、晶粒特征、析出相或缺陷迹象，"
        "判断组织状态是否正常。"
    ),
    "chemical": (
        "阅读{section_title}：{slot_labels}，"
        "概述主要元素组成特征，"
        "判断各元素含量是否符合标准要求。"
    ),
    "mechanical": (
        "阅读{section_title}：{slot_labels}，"
        "概述测试数据分布是否均匀及是否存在明显偏离，"
        "判断性能指标是否满足要求。"
    ),
    "generic": (
        "观察{section_title}中的{slot_labels}，"
        "客观描述图片中可见的关键特征与异常。"
    ),
}

# 引导文字模板（动态占位符，无硬编码样品名）
INTRO_TYPE_TEMPLATES = {
    "overview": "{background_summary}，具体位置如{fig_label}所示。",
    "macro": "在体视显微镜下进行{section_title_clean}，具体形貌如{fig_label}所示。",
    "sem": "在扫描电镜下进行{section_title_clean}，具体分析如{fig_label}所示。",
    "metallography": "进行{section_title_clean}，其结果如{fig_label}所示。",
    "chemical": "进行{section_title_clean}，具体试验结果如{fig_label}所示。",
    "mechanical": "进行{section_title_clean}，具体试验结果如{fig_label}所示。",
    "generic": "进行{section_title_clean}，具体结果如{fig_label}所示。",
}

PLACEHOLDER_MAP = {
    "sec_1": PH_BASIC,
    "sec_2": PH_MACRO,
    "sec_3": PH_MICRO,
    "sec_4": PH_MICROSTRUCT,
    "sec_5": PH_CHEM,
    "sec_6": PH_PERF,
}


# -------------------------
# 基础工具
# -------------------------
def ensure_dir(p: Path) -> Path:
    p.mkdir(parents=True, exist_ok=True)
    return p


def normalize_text(val: Any) -> str:
    if val is None:
        return "none"
    s = str(val).strip()
    return s if s else "none"


def safe_name(s: str, fallback: str = "project") -> str:
    s = (s or "").strip()
    if not s:
        return fallback
    s = re.sub(r"[^\w\u4e00-\u9fff\- ]+", "_", s)
    s = re.sub(r"\s+", "_", s).strip("_")
    return s or fallback


def find_template_docx() -> Path:
    script_dir = Path(__file__).resolve().parent
    local_desktop = Path.home() / "Desktop"
    icloud_desktop = Path.home() / "Library" / "Mobile Documents" / "com~apple~CloudDocs" / "Desktop"
    candidates = []
    for base in [script_dir, local_desktop, icloud_desktop]:
        for name in ["模版.docx", "模板.docx"]:
            candidates.append(base / name)
    for p in candidates:
        if p.exists():
            return p
    raise FileNotFoundError("找不到模板文件，已尝试：\n" + "\n".join(str(x) for x in candidates))


def find_template_docx_cover() -> Path:
    """查找精简封面模板（模版_cover.docx），找不到则回退到完整模板。"""
    script_dir = Path(__file__).resolve().parent
    local_desktop = Path.home() / "Desktop"
    icloud_desktop = Path.home() / "Library" / "Mobile Documents" / "com~apple~CloudDocs" / "Desktop"
    for base in [script_dir, local_desktop, icloud_desktop]:
        p = base / "模版_cover.docx"
        if p.exists():
            return p
    # 回退到完整模板
    print("WARNING: 模版_cover.docx not found, falling back to 模版.docx")
    return find_template_docx()


def set_run_font(run, font_name: str, size_pt: float) -> None:
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    r = run._element.rPr
    if r is None:
        r = run._element.get_or_add_rPr()
    rFonts = r.rFonts
    if rFonts is None:
        rFonts = r.get_or_add_rFonts()
    rFonts.set(qn("w:eastAsia"), font_name)
    rFonts.set(qn("w:ascii"), font_name)
    rFonts.set(qn("w:hAnsi"), font_name)


# -------------------------
# 图片处理
# -------------------------
def to_rgb_image(src_path: Path) -> Image.Image:
    im = Image.open(src_path)
    im = ImageOps.exif_transpose(im)
    if im.mode == "RGBA":
        bg = Image.new("RGB", im.size, (255, 255, 255))
        bg.paste(im, mask=im.split()[-1])
        im = bg
    elif im.mode != "RGB":
        im = im.convert("RGB")
    return im


def aspect_fit_pad(im: Image.Image, target_size: Tuple[int, int]) -> Tuple[Image.Image, Dict[str, Any]]:
    tw, th = target_size
    ow, oh = im.size
    scale = min(tw / ow, th / oh)
    nw = max(1, int(round(ow * scale)))
    nh = max(1, int(round(oh * scale)))
    im_resized = im.resize((nw, nh), Image.Resampling.LANCZOS)
    canvas = Image.new("RGB", (tw, th), (255, 255, 255))
    x = (tw - nw) // 2
    y = (th - nh) // 2
    canvas.paste(im_resized, (x, y))
    info = {
        "original_px": {"w": ow, "h": oh},
        "resized_px": {"w": nw, "h": nh},
        "final_px": {"w": tw, "h": th},
        "method": "aspect-fit + white padding",
    }
    return canvas, info


def image_to_jpg_bytes(im: Image.Image, quality: int = 95) -> bytes:
    buf = io.BytesIO()
    im.save(buf, format="JPEG", quality=quality, optimize=True)
    return buf.getvalue()


def jpg_bytes_to_base64(data: bytes) -> str:
    return base64.b64encode(data).decode("utf-8")


def data_url_from_b64_jpg(b64: str) -> str:
    return "data:image/jpeg;base64," + b64


# -------------------------
# OpenAI 兼容调用
# -------------------------
def call_openai_chat(
    model: str,
    messages: List[Dict[str, Any]],
    temperature: float = 0.2,
    max_tokens: int = 800,
    timeout: int = 180,
) -> str:
    url = BASE_URL.rstrip("/") + "/chat/completions"
    headers = {"Authorization": f"Bearer {API_KEY}", "Content-Type": "application/json"}
    payload = {"model": model, "messages": messages, "temperature": temperature, "max_tokens": max_tokens}
    r = requests.post(url, headers=headers, json=payload, timeout=timeout)
    r.raise_for_status()
    data = r.json()
    return data["choices"][0]["message"]["content"]


# -------------------------
# 知识库 API 调用
# -------------------------
def kb_login() -> str:
    """调用 KB-Matrix 登录接口，返回 accessToken；失败返回空字符串。"""
    if not KB_AUTH_URL or not KB_USERNAME:
        return ""
    url = KB_AUTH_URL.rstrip("/") + "/auth/login-local"
    payload = {"username": KB_USERNAME, "password": KB_PASSWORD}
    try:
        r = requests.post(url, json=payload, timeout=10)
        r.raise_for_status()
        data = r.json()
        return (data.get("data") or {}).get("accessToken", "")
    except Exception as e:
        print(f"KB login failed: {e}")
        return ""


def kb_list(token: str) -> List[Dict[str, Any]]:
    """调用知识库列表接口，返回 [{"id": ..., "name": ...}, ...]。"""
    if not KB_API_URL or not token:
        return []
    url = KB_API_URL.rstrip("/") + "/kb-base/list"
    headers = {"Authorization": token}
    try:
        r = requests.get(url, headers=headers, timeout=10)
        r.raise_for_status()
        return (r.json().get("data") or [])
    except Exception as e:
        print(f"KB list failed: {e}")
        return []


def kb_retrieve(token: str, knowledge_id: str, query: str, top_k: int = 5, score_threshold: float = 0.5) -> str:
    """调用知识库检索接口，拼接所有命中文本返回；失败返回空字符串。"""
    if not KB_API_URL or not token or not knowledge_id:
        return ""
    url = KB_API_URL.rstrip("/") + "/kb-search/retrieval"
    headers = {"Authorization": token, "Content-Type": "application/json"}
    payload = {
        "knowledge_id": knowledge_id,
        "query": query,
        "retrieval_setting": {"top_k": top_k, "score_threshold": score_threshold},
    }
    try:
        r = requests.post(url, headers=headers, json=payload, timeout=30)
        r.raise_for_status()
        records = (r.json().get("records") or [])
        parts = []
        for rec in records:
            content = rec.get("content", "").strip()
            if content:
                parts.append(content)
        return "\n\n".join(parts)
    except Exception as e:
        print(f"KB retrieve failed: {e}")
        return ""


# -------------------------
# Word 文本替换（替换进去的文字统一宋体五号）
# -------------------------
def _replace_in_paragraph_with_style(paragraph, mapping: Dict[str, str], font_name: str, size_pt: float) -> bool:
    full_text = "".join(run.text for run in paragraph.runs)
    if not full_text: return False
    
    modified = False
    new_text = full_text
    
    for k, v in mapping.items():
        if k in new_text:
            new_text = new_text.replace(k, v)
            modified = True
            print(f"DEBUG: Text replaced '{k}' -> '{v}'")

    if modified:
        # 重建段落：清空所有旧的 run，保留样式放入第一个 run 或新加的 run
        for run in paragraph.runs:
            run.text = ""
        if paragraph.runs:
            run0 = paragraph.runs[0]
        else:
            run0 = paragraph.add_run()
        run0.text = new_text
        set_run_font(run0, font_name, size_pt)
    return modified


# -------------------------


def replace_text_placeholders_with_style(doc: Document, mapping: Dict[str, str], font_name: str, size_pt: float, exact_match: bool = False) -> None:
    for p in doc.paragraphs:
        _replace_in_paragraph_with_style(p, mapping, font_name, size_pt)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    _replace_in_paragraph_with_style(p, mapping, font_name, size_pt)


# -------------------------
# Word 图片替换
# -------------------------
def _add_image_slot_layout(doc, anchor_p, image_path: Path, mark: str, label: str, width_cm: float):
    """
    在 anchor_p 之后插入三层布局：图片、标记 (a)、描述文字。
    如果 anchor_p 为 None，则在文档末尾追加。
    """
    from docx.oxml import OxmlElement
    
    def insert_p_after(target_p, text="", is_bold=False, is_center=True, font_size=10.5):
        new_p = doc.add_paragraph()
        if target_p:
            target_p._p.addnext(new_p._p)
        run = new_p.add_run(text)
        run.bold = is_bold
        set_run_font(run, FONT_NAME_CN, font_size)
        if is_center:
            new_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        return new_p

    # 注意顺序：先插底层，再插中层，最后插顶层（因为是 addnext，后插的会在目标段落的紧下方）
    # 最终期望：[图片] -> [标记] -> [描述]
    # 我们按照 描述 -> 标记 -> 图片 的顺序 addnext 到 anchor_p 之后
    
    # 1. 描述段落
    p_desc = insert_p_after(anchor_p, label, is_bold=False)
    # 2. 标记段落
    p_mark = insert_p_after(anchor_p, f"({mark})", is_bold=True)
    # 3. 图片段落
    p_img = insert_p_after(anchor_p)
    run_img = p_img.add_run()
    run_img.add_picture(str(image_path), width=Cm(width_cm))
    
    return p_desc # 返回最后一项作为新的锚点

def _insert_image_into_placeholder_paragraph(paragraph, image_path: Path, mark: str, label: str, width_cm: float) -> bool:
    """
    替换模板中的占位符段落为三层布局。
    由于占位符段落本身就在文档中，我们直接在该段落及其后方展开。
    """
    # 清空原段落内容，作为第一层：图片
    for run in paragraph.runs:
        run.text = ""
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=Cm(width_cm))
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 在其后方插入 标记 和 描述
    from docx.oxml import OxmlElement
    
    # 描述
    p_desc = doc_add_p_after(paragraph)
    r_desc = p_desc.add_run(label)
    set_run_font(r_desc, FONT_NAME_CN, 10.5)
    p_desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 标记 (插在图片和描述之间)
    p_mark = doc_add_p_after(paragraph)
    r_mark = p_mark.add_run(f"({mark})")
    r_mark.bold = True
    set_run_font(r_mark, FONT_NAME_CN, 10.5)
    p_mark.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    return p_desc # 返回最底部的段落

def doc_add_p_after(p):
    new_p = p.__class__(p._p.getparent().add_p(), p._parent) # 这种方式在 python-docx 中比较 hack
    p._p.addnext(new_p._p)
    return new_p


def replace_image_placeholders_v2(doc: Document, placeholder_to_info: Dict[str, Dict[str, Any]], width_cm: float) -> Dict[str, Any]:
    """
    返回 slot_id -> 最后一层段落对象的映射，用于后续锚点定位。
    placeholder_to_info: { "{{image_1.1}}": {"path": Path, "mark": "a", "label": "..."} }
    """
    pattern = re.compile(r"\{\{image_[^}]+\}\}")
    slot_anchors = {} # slot_id -> last_p
    
    def process_p(p):
        text = p.text
        if "{{" not in text or "image_" not in text: return
        matches = pattern.findall(text)
        for m in matches:
            if m in placeholder_to_info:
                info = placeholder_to_info[m]
                print(f"DEBUG: Replacing template placeholder {m} with 3-layer layout")
                # 记录 slot_id
                slot_id = m.replace("{{image_", "").replace("}}", "")
                last_p = _replace_placeholder_with_3layers(doc, p, m, info["path"], info["mark"], info["label"], width_cm)
                slot_anchors[slot_id] = last_p
            else:
                p.text = p.text.replace(m, "")

    for p in list(doc.paragraphs): process_p(p)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in list(cell.paragraphs): process_p(p)
    return slot_anchors

def _replace_placeholder_with_3layers(doc, p, placeholder, img_path, mark, label, width_cm):
    # 1. 图片层：替换当前段落的占位符内容
    p.text = p.text.replace(placeholder, "")
    run = p.add_run()
    run.add_picture(str(img_path), width=Cm(width_cm))
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 2. 标记层 & 描述层
    # 采用 addnext 方式，先插底层（描述），再插中层（标记）
    p_desc = doc.add_paragraph()
    p._p.addnext(p_desc._p)
    r_desc = p_desc.add_run(label)
    set_run_font(r_desc, FONT_NAME_CN, 10.5)
    p_desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p_mark = doc.add_paragraph()
    p._p.addnext(p_mark._p)
    r_mark = p_mark.add_run(f"({mark})")
    r_mark.bold = True
    set_run_font(r_mark, FONT_NAME_CN, 10.5)
    p_mark.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    return p_desc

def _add_image_slot_layout_v2(doc, anchor, image_path, mark, label, width_cm):
    """在 anchor 段落之后按顺序插入图片、标记、描述三层段落。"""
    # 描述层 (最下方)
    p_desc = doc.add_paragraph()
    anchor._p.addnext(p_desc._p)
    r_desc = p_desc.add_run(label)
    set_run_font(r_desc, FONT_NAME_CN, 10.5)
    p_desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 标记层 (中间)
    p_mark = doc.add_paragraph()
    anchor._p.addnext(p_mark._p)
    r_mark = p_mark.add_run(f"({mark})")
    r_mark.bold = True
    set_run_font(r_mark, FONT_NAME_CN, 10.5)
    p_mark.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 图片层 (最上方，紧跟 anchor)
    p_img = doc.add_paragraph()
    anchor._p.addnext(p_img._p)
    r_img = p_img.add_run()
    r_img.add_picture(str(image_path), width=Cm(width_cm))
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    return p_desc # 返回描述段落作为后续的插入锚点


# -------------------------
# VLM Prompt：更短 + 允许标点 + 输出更精简（2~4 句）
# -------------------------
def build_vlm_prompt(background_overview: str, what: str, figure_label: str = "") -> str:
    return (
        "你是一名材料失效分析工程师。\n"
        f"背景概述：{background_overview}\n"
        f"任务：{what}\n"
        "要求：只根据图片做客观描述与对比，不要编造，不要引用外部资料，不要下结论。\n"
        "输出格式：以'由{fig_label}可知，'或'{fig_label}显示了...'开头，"
        "然后用2到4句话客观描述图片中的关键特征。\n"
        f"图号引用：{figure_label}\n"
    )


def build_vlm_messages(prompt: str, image_b64_list: List[str]) -> List[Dict[str, Any]]:
    content = [{"type": "text", "text": prompt}]
    for b64 in image_b64_list:
        content.append({"type": "image_url", "image_url": {"url": data_url_from_b64_jpg(b64)}})
    return [{"role": "user", "content": content}]


# -------------------------
# LLM 最终生成：分析与讨论/结论/报告摘要/关键词（JSON）
# -------------------------
def build_final_llm_prompt(
    background_overview: str,
    analyses: list,
    kb_knowledge: str = "",
) -> str:
    kb_section = ""
    if kb_knowledge and kb_knowledge.strip():
        kb_section = f"\n相关知识库内容：\n{kb_knowledge}\n"

    # 动态拼接所有分析结果
    analysis_parts = []
    for a in analyses:
        title = a.get("title", "")
        result = a.get("result", "none")
        analysis_parts.append(f"\n{title}：\n{result}\n")
    analysis_text = "".join(analysis_parts)

    return (
        "你是一名材料失效分析报告撰写专家。\n"
        "请基于背景概述与各部分图片分析内容，生成报告中的关键章节。\n"
        "必须输出严格的 JSON，不要包含多余文字，不要用 markdown 代码块包裹。\n"
        "JSON keys 必须是：分析与讨论 结论 报告摘要 关键词\n"
        "要求：\n"
        "1 分析与讨论：2到3段，综合所有分析结果，语言专业但不过长，引用图号时用'图X'格式。\n"
        "2 结论：先写一段概述（以'综上所述'开头，总结样品信息、工况、失效模式，约80字），"
        "然后分条列出结论，每条以（1）（2）（3）开头，每条20到50字。"
        "结论应是简洁判定而非重复分析细节。"
        "第1条：断裂模式判定。"
        "第2条：材料/性能是否达标判定。"
        "第3条：具体可操作的改进建议（引用部件名称）。\n"
        "3 报告摘要：一段话，80到150字。\n"
        "4 关键词：4到8个词，用中文顿号分隔。\n"
        "\n背景概述：\n"
        f"{background_overview}\n"
        f"{analysis_text}"
        f"{kb_section}"
    )


def parse_llm_json(text: str) -> Dict[str, str]:
    def normalize_section(val: Any) -> str:
        if isinstance(val, list):
            items = [str(x).strip() for x in val if str(x).strip()]
            cleaned = []
            summary_prefix = ""
            for i, item in enumerate(items):
                # 去除已有的编号前缀，如 (1)、（1）、1.、1、等
                stripped = re.sub(r'^[（\(]\d+[）\)]\s*', '', item)
                stripped = re.sub(r'^\d+[.、]\s*', '', stripped)
                # 概述段落（以"综上所述"开头）不加编号
                if i == 0 and stripped.startswith("综上所述"):
                    summary_prefix = stripped
                    continue
                cleaned.append(f"（{len(cleaned) + 1}）{stripped}")
            result = "\n".join(cleaned)
            if summary_prefix:
                result = summary_prefix + "\n" + result
            return result if result else ""
        if isinstance(val, dict):
            lines = []
            for k, v in val.items():
                ks = str(k).strip()
                vs = str(v).strip()
                if ks and vs:
                    lines.append(f"{ks}: {vs}")
            return "\n".join(lines)
        return str(val or "").strip()

    if not text:
        return {}
    s = text.strip()
    # 去除可能的 markdown 代码块包裹
    if s.startswith("```"):
        s = re.sub(r"^```(?:json)?\s*", "", s)
        s = re.sub(r"\s*```\s*$", "", s)
    m = re.search(r"\{.*\}", s, flags=re.S)
    if m:
        s = m.group(0)
    try:
        obj = json.loads(s)
        return {
            "分析与讨论": normalize_section(obj.get("分析与讨论", "")),
            "结论": normalize_section(obj.get("结论", "")),
            "报告摘要": normalize_section(obj.get("报告摘要", "")),
            "关键词": normalize_section(obj.get("关键词", "")),
        }
    except Exception as e:
        print(f"WARNING: JSON parse failed: {e}")
        print(f"WARNING: Raw text was: {s[:300]}")
        return {}


# -------------------------
# 主流程（Flask SSE 流式版）
# -------------------------
_UPLOAD_TMP = Path(tempfile.gettempdir()) / "report_tool_uploads"
_UPLOAD_TMP.mkdir(parents=True, exist_ok=True)

LLM_CHOICES = [
    "Qwen/Qwen3-32B",
    "deepseek-ai/DeepSeek-R1-Distill-Qwen-32B",
    "Qwen/Qwen2.5-72B-Instruct",
]
VLM_CHOICES = [
    "Qwen/Qwen3-VL-32B-Instruct",
    "Qwen/Qwen2.5-VL-72B-Instruct",
    "Pro/Qwen/Qwen2.5-VL-7B-Instruct",
]
KB_CHOICES = [
    "工艺知识库",
    "材料知识库",
    "检测标准知识库",
]


# ── 封面页数插入 ──

def _get_page_count_win32(docx_path: str) -> int | None:
    """使用 Word COM 获取文档总页数（仅 Windows）。"""
    try:
        import win32com.client
        import pythoncom
        pythoncom.CoInitialize()
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        abs_path = os.path.abspath(docx_path)
        doc = word.Documents.Open(abs_path)
        count = doc.ComputeStatistics(2)  # wdStatisticPages = 2
        doc.Close(False)
        word.Quit()
        pythoncom.CoUninitialize()
        return count
    except Exception as e:
        print(f"WARNING: page count via win32com failed: {e}")
        return None


def _insert_page_count(doc: Document, page_count: int) -> None:
    """在封面 '编号：' 段落后插入 '页数：X' 段落。"""
    for p in doc.paragraphs:
        if "编号" in p.text and ("：" in p.text or ":" in p.text):
            # 提取与"编号："相同的行首空格
            idx = p.text.index("编")
            prefix = p.text[:idx]
            # 在"编号："段落后插入新段落
            new_p = doc.add_paragraph()
            p._p.addnext(new_p._p)
            new_p.alignment = p.alignment
            run = new_p.add_run(f"{prefix}页数：{page_count}")
            # 复制字体设置
            run.font.name = "宋体"
            run.font.size = Pt(10.5)
            r = run._element.get_or_add_rPr()
            rFonts = r.find(qn("w:rFonts"))
            if rFonts is None:
                from docx.oxml import OxmlElement
                rFonts = OxmlElement("w:rFonts")
                r.append(rFonts)
            rFonts.set(qn("w:eastAsia"), "宋体")
            rFonts.set(qn("w:ascii"), "宋体")
            rFonts.set(qn("w:hAnsi"), "宋体")
            break


def _generate_report_stream(form, files):
    """Generator: yields SSE data strings for Flask streaming response."""

    def sse(msg: str, file_path: str | None = None, done: bool = False) -> str:
        print(msg)
        payload = {"msg": msg, "file": file_path, "done": done}
        return f"data: {json.dumps(payload, ensure_ascii=False)}\n\n"

    try:
        yield sse("⏳ 正在检查模板文件...")
        if USE_DYNAMIC_ASSEMBLY:
            template_path = find_template_docx_cover()
        else:
            template_path = find_template_docx()

        project_name_val = form.get("项目名称", "")
        background_overview = form.get("背景概述", "") or "none"
        llm_model = form.get("llm_model", "") or LLM_MODEL
        vlm_model = form.get("vlm_model", "") or VLM_MODEL
        kb_id = form.get("kb_id", "")  # 知识库 ID（前端下拉框传来的值）
        
        # 解析动态配置
        config_str = form.get("image_config", "[]")
        image_config = json.loads(config_str)

        text_data = {
            "项目名称": normalize_text(project_name_val),
            "背景概述": normalize_text(background_overview),
        }

        ts = datetime.now().strftime("%Y%m%d_%H%M%S")
        proj = safe_name(text_data["项目名称"], fallback="project")
        job_id = f"{ts}_{proj}"
        job_dir = ensure_dir(JOB_ROOT / job_id)
        images_dir = ensure_dir(job_dir / "images_jpg")

        images_payload: Dict[str, Any] = {}
        placeholder_to_image_path: Dict[str, Path] = {}
        slot_to_b64: Dict[str, str] = {}
        total_images = 0
        total_bytes = 0

        yield sse("⏳ 正在处理上传图片...")
        
        # 展平所有槽位以便处理
        all_slots = []
        for section in image_config:
            for slot in section.get("slots", []):
                all_slots.append((section, slot))

        total_slots_count = sum(1 for _, slot in all_slots if files.get(f"image_{slot['id']}"))
        upload_idx = 0
        for idx, (section, slot) in enumerate(all_slots, start=1):
            slot_id = slot["id"]
            display = slot["label"]
            category = section["title"]
            # 兼容旧逻辑：1-4 固宽，5-6 原样
            section_num = section["title"].split(".")[0].strip() if section["title"] else ""
            fixed_size = section_num in FIXED_SIZE_SECTIONS
            # 自动映射占位符，例如 id 是 1.1 则映射到 {{image_1.1}}
            placeholder = f"{{{{image_{slot_id}}}}}"

            # 从预存的文件字典中获取路径
            tmp_path = files.get(f"image_{slot_id}")
            if not tmp_path:
                images_payload[slot_id] = {
                    "display": display, "category": category, "provided": False,
                    "placeholder": placeholder, "jpg_relpath": None,
                }
                continue

            try:
                im = to_rgb_image(tmp_path)
                process_info = {
                    "uploaded_filename": tmp_path.name,
                    "rule": "fixed_size" if fixed_size else "original_size",
                    "original_px": {"w": im.size[0], "h": im.size[1]},
                }
                if fixed_size:
                    im, fit_info = aspect_fit_pad(im, TARGET_SIZE)
                    process_info.update(fit_info)
                else:
                    process_info.update({
                        "final_px": {"w": im.size[0], "h": im.size[1]},
                        "method": "convert-to-RGB-only",
                    })

                jpg_bytes = image_to_jpg_bytes(im, quality=95)
                b64 = jpg_bytes_to_base64(jpg_bytes)
                jpg_name = f"{slot_id.replace('.', '_')}.jpg"
                jpg_path = images_dir / jpg_name
                jpg_path.write_bytes(jpg_bytes)
                total_images += 1
                total_bytes += len(jpg_bytes)
                
                images_payload[slot_id] = {
                    "display": display, "category": category, "provided": True,
                    "slot_id": slot_id,
                    "jpg_relpath": f"images_jpg/{jpg_name}",
                    "jpg_base64": b64,
                }
                # 使用 slot_id 作为键
                placeholder_to_image_path[slot_id] = jpg_path
                slot_to_b64[slot_id] = b64
            finally:
                if tmp_path.exists(): tmp_path.unlink()
            upload_idx += 1
            yield sse(f"⏳ 图片处理进度：{upload_idx}/{total_slots_count}")

        # 生成 Payload
        payload_data = {
            "job_id": job_id, "created_at": datetime.now().isoformat(),
            "text": text_data, "images": images_payload,
        }
        (job_dir / "payload.json").write_text(json.dumps(payload_data, ensure_ascii=False, indent=2), encoding="utf-8")

        report_name = f"报告_{proj}_{ts}.docx"
        report_path = job_dir / report_name
        shutil.copyfile(template_path, report_path)
        doc = Document(str(report_path))

        yield sse("⏳ 正在替换模板文本...")
        base_mapping = {placeholder: text_data[name] for name, placeholder in TEXT_FIELDS}
        for ph in OVERVIEW_SYNC_PLACEHOLDERS:
            base_mapping[ph] = background_overview
        replace_text_placeholders_with_style(doc, base_mapping, FONT_NAME_CN, FONT_SIZE_5_PT)

        if not API_KEY.strip():
            doc.save(str(report_path))
            yield sse("✅ 生成完成 (API_KEY 为空，跳过模型分析)", file_path=str(report_path), done=True)
            return

        # -------------------------
        # 动态 VLM 分析
        # -------------------------
        # ── 判断章节是否含表格：基于 slot label 含"表"字 ──
        def _section_has_table(section: dict) -> bool:
            return any("表" in s.get("label", "") for s in section.get("slots", []))

        # 预计算每个 section 的图号标签（用于 VLM prompt）
        _fig_count = 0
        _tbl_count = 0
        section_fig_labels: Dict[str, str] = {}
        for _sec_idx, _sec in enumerate(image_config):
            _sid = _sec.get("id", "")
            _slots = [s for s in _sec.get("slots", []) if s["id"] in slot_to_b64]
            if not _slots:
                continue
            if _section_has_table(_sec):
                _tbl_count += 1
                section_fig_labels[_sid] = f"表{_tbl_count}"
            else:
                _fig_count += 1
                section_fig_labels[_sid] = f"图{_fig_count}"

        # ── 表格图片提取提示词 ──
        TABLE_EXTRACTION_PROMPT = (
            "你是一个表格数据提取专家。请判断这张图片是否为数据表格（包含行列结构的文字数据）。\n"
            "如果是表格，提取所有单元格内容，严格按以下JSON格式返回（不要markdown标记）：\n"
            '{"is_table": true, "headers": ["列标题1", "列标题2"], '
            '"rows": [["行1值1", "行1值2"]]}\n'
            "如果不是表格（如照片、图谱、曲线图等），返回：\n"
            '{"is_table": false}\n'
            "注意：保留原始数值精度和单位，空单元格返回空字符串。只返回JSON。"
        )

        def extract_table_data(image_b64: str) -> dict | None:
            """尝试从图片中提取表格数据。返回 {"headers": [...], "rows": [[...]]} 或 None。"""
            messages = build_vlm_messages(TABLE_EXTRACTION_PROMPT, [image_b64])
            try:
                out = call_openai_chat(
                    model=vlm_model, messages=messages,
                    temperature=0.1, max_tokens=500)
                text = out.strip()
                if text.startswith("```"):
                    text = text.removeprefix("```json").removeprefix("```").removesuffix("```")
                data = json.loads(text)
                if data.get("is_table") and data.get("headers") and data.get("rows"):
                    return {"headers": data["headers"], "rows": data["rows"]}
            except (json.JSONDecodeError, KeyError, Exception) as e:
                print(f"WARNING: table extraction failed: {e}")
            return None

        def run_dynamic_vlm(section, is_first: bool = False):
            s_id = section.get("id", "")
            s_title = section.get("title", "")
            slots = [s for s in section.get("slots", []) if s["id"] in slot_to_b64]
            if not slots: return "none"

            slot_labels = "、".join([s["label"] for s in slots])
            sec_type = detect_section_type(s_title, is_first=is_first)
            base_task = VLM_TYPE_PROMPTS.get(
                sec_type, VLM_TYPE_PROMPTS["generic"])
            task_desc = base_task.format(
                section_title=s_title, slot_labels=slot_labels)
            fig_label = section_fig_labels.get(s_id, "")

            b64s = [slot_to_b64[s["id"]] for s in slots]
            prompt = build_vlm_prompt(background_overview, task_desc, figure_label=fig_label)
            messages = build_vlm_messages(prompt, b64s)
            try:
                out = call_openai_chat(model=vlm_model, messages=messages, temperature=0.2, max_tokens=500)
                return out.strip() if out else "none"
            except Exception as e:
                print(f"WARNING: VLM call failed for {s_id}: {e}")
                return "none"

        yield sse("⏳ 正在调用 VLM（动态图片分析）...")
        analysis_results = {}
        processed_sec_ids = set()
        for sec_idx, section in enumerate(image_config):
            sec_id = section.get("id", "")
            yield sse(f"⏳ 正在分析章节：{section.get('title', sec_id)}")
            res = run_dynamic_vlm(section, is_first=(sec_idx == 0))
            analysis_results[sec_id] = res
            processed_sec_ids.add(sec_id)

        # ── 表格图片检测与提取（slot label 含"表"字） ──
        yield sse("⏳ 正在检测表格图片...")
        slot_table_data: Dict[str, dict] = {}
        for _sec in image_config:
            for _slot in _sec.get("slots", []):
                if "表" not in _slot.get("label", ""):
                    continue
                _s_id = _slot["id"]
                if _s_id not in slot_to_b64:
                    continue
                td = extract_table_data(slot_to_b64[_s_id])
                if td:
                    slot_table_data[_s_id] = td
                    yield sse(f"  ✅ 检测到表格：{_slot['label']}")

        # -------------------------
        # 章节构建
        # -------------------------
        if USE_DYNAMIC_ASSEMBLY:
            # ---- 新方案：程序化构建所有内容章节 ----
            yield sse("⏳ 正在构建章节内容...")

            # 1. 收集每个章节的图片槽位
            section_to_slots_with_images: Dict[str, List[dict]] = {}
            for section in image_config:
                sec_id = section["id"]
                slots_with_img = []
                for slot in section.get("slots", []):
                    s_id = slot["id"]
                    if s_id in placeholder_to_image_path:
                        slots_with_img.append({
                            "id": s_id,
                            "label": slot["label"],
                            "path": placeholder_to_image_path[s_id],
                            "table_data": slot_table_data.get(s_id),
                        })
                section_to_slots_with_images[sec_id] = slots_with_img

            # 2. 找到插入锚点："3 分析与讨论" 标题之前的段落
            insert_anchor = None
            for i, p in enumerate(doc.paragraphs):
                text = p.text.strip()
                # 匹配 "3 分析与讨论" 或 "3  分析与讨论" 等标题行（非占位符）
                if text and "分析" in text and "讨论" in text and not text.startswith("{{"):
                    # 使用其前一个段落作为锚点
                    if i > 0:
                        insert_anchor = doc.paragraphs[i - 1]
                    else:
                        insert_anchor = p
                    break

            if not insert_anchor:
                # 回退：找到 {{分析与讨论}} 占位符的前前段落
                for i, p in enumerate(doc.paragraphs):
                    if "{{" in p.text and "分析与讨论" in p.text:
                        insert_anchor = doc.paragraphs[max(0, i - 2)]
                        break

            if not insert_anchor:
                insert_anchor = doc.paragraphs[-1] if doc.paragraphs else None

            current_anchor = insert_anchor

            # 3. 插入项目名称标题（居中、宋体三号加粗）
            from section_builders import set_run_font as _set_run_font
            proj_title_p = doc.add_paragraph()
            current_anchor._p.addnext(proj_title_p._p)
            proj_title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            proj_title_run = proj_title_p.add_run(project_name_val)
            proj_title_run.bold = True
            _set_run_font(proj_title_run, "宋体", 16.0)  # 三号 = 16pt
            current_anchor = proj_title_p

            # 4. 按顺序构建所有章节
            #    第一个手风琴 → "1 概述"
            #    其余手风琴 → "2 分析过程及结果"下的子章节 2.1、2.2...
            figure_counter = 0  # 全局连续图编号
            table_counter = 0   # 表格编号（化学分析、性能验证等）
            sub_section_idx = 0  # 分析过程子章节计数（2.1, 2.2...）
            analysis_header_added = False

            # 截取背景概述前 80 字用于引导文字
            bg_summary = (background_overview[:80] if background_overview and background_overview != "none"
                          else "")

            for sec_idx, section in enumerate(image_config):
                sec_id = section["id"]
                slots_with_img = section_to_slots_with_images.get(sec_id, [])

                # 跳过无图片的空章节
                if not slots_with_img:
                    continue

                analysis_text = analysis_results.get(sec_id, "none")
                group_title = section.get("groupTitle", "")
                raw_title = section.get("title", "")

                # 按内容判断是图还是表（基于 slot label 含"表"字）
                is_table = any("表" in s.get("label", "") for s in slots_with_img)

                # 按标题关键词检测类型（用于引导文字模板）
                sec_type = detect_section_type(
                    raw_title, is_first=(sec_idx == 0))

                if is_table:
                    table_counter += 1
                    fig_num = table_counter
                    fig_label = f"表{fig_num}"
                else:
                    figure_counter += 1
                    fig_num = figure_counter
                    fig_label = f"图{fig_num}"

                if sec_idx == 0:
                    # 第一个手风琴 → "1 概述"
                    section_title = "1 概述"
                    figure_caption = f"图{fig_num} {group_title}" if group_title else ""
                else:
                    if not analysis_header_added:
                        # 在第一个分析子章节前插入"2 分析过程及结果"大标题
                        current_anchor = build_section_title(
                            doc, current_anchor, "2 分析过程及结果")
                        analysis_header_added = True

                    sub_section_idx += 1
                    # 避免编号重复：如果 title 已含数字前缀（如 "2.1 宏观断口分析"），直接使用
                    if raw_title and raw_title[0].isdigit() and '.' in raw_title.split()[0]:
                        section_title = raw_title
                    else:
                        section_title = f"2.{sub_section_idx} {raw_title}"
                    figure_caption = f"{fig_label} {group_title}" if group_title else ""

                # 生成引导文字（使用类型驱动的模板）
                intro_text = ""
                intro_tpl = INTRO_TYPE_TEMPLATES.get(sec_type, "")
                if intro_tpl:
                    section_title_clean = re.sub(
                        r"^[\d.]+\s*", "", raw_title)
                    intro_text = intro_tpl.format(
                        background_summary=bg_summary,
                        section_title_clean=section_title_clean,
                        fig_label=fig_label,
                    )

                current_anchor = build_complete_section(
                    doc, current_anchor,
                    section_title=section_title,
                    image_slots=slots_with_img,
                    analysis_text=analysis_text,
                    figure_caption=figure_caption,
                    intro_text=intro_text,
                )

        else:
            # ---- 旧方案：模板占位符替换 + 溢出插入 ----
            print(f"DEBUG: Starting report sync logic (legacy). image_config sections: {len(image_config)}")

            title_mapping = {}
            vlm_text_mapping = {}
            section_to_slots_with_images = {}

            for section in image_config:
                sec_id = section["id"]
                slots_with_img = []
                for slot in section.get("slots", []):
                    s_id = slot["id"]
                    if s_id in placeholder_to_image_path:
                        slots_with_img.append({
                            "id": s_id,
                            "label": slot["label"],
                            "path": placeholder_to_image_path[s_id]
                        })
                section_to_slots_with_images[sec_id] = slots_with_img

            title_mapping = {}
            for section in image_config:
                sec_id = section["id"]
                if sec_id.startswith("sec_"):
                    num = sec_id[4:]
                    title_mapping[f"{{{{sec_title_{num}}}}}"] = section["title"]

            vlm_text_mapping = {}
            for sec_id, ph in PLACEHOLDER_MAP.items():
                current_section = next((x for x in image_config if x["id"] == sec_id), None)
                if current_section:
                    vlm_text_mapping[ph] = analysis_results.get(sec_id, "none")
                else:
                    vlm_text_mapping[ph] = ""

            replace_text_placeholders_with_style(doc, title_mapping, FONT_NAME_CN, FONT_SIZE_5_PT)
            replace_text_placeholders_with_style(doc, vlm_text_mapping, FONT_NAME_CN, FONT_SIZE_5_PT)

            placeholder_to_info = {}
            for sec_id, slots in section_to_slots_with_images.items():
                for idx, slot_info in enumerate(slots):
                    ph = f"{{{{image_{slot_info['id']}}}}}"
                    mark = chr(ord('a') + idx)
                    placeholder_to_info[ph] = {
                        "path": slot_info["path"],
                        "mark": mark,
                        "label": slot_info["label"]
                    }

            slot_anchors = replace_image_placeholders_v2(doc, placeholder_to_info, width_cm=WORD_IMAGE_WIDTH_CM)

            section_anchors = {}
            for section in image_config:
                sec_id = section["id"]
                slots = section_to_slots_with_images.get(sec_id, [])
                last_p = None
                for s in slots:
                    if s["id"] in slot_anchors:
                        last_p = slot_anchors[s["id"]]
                if last_p:
                    section_anchors[sec_id] = last_p

            yield sse("⏳ 正在对齐章节结构并优化排版...")

            prev_sec_anchor = None
            for p in doc.paragraphs:
                if "{{背景概述}}" in p.text or "{{概述}}" in p.text:
                    prev_sec_anchor = p
                    break
            if not prev_sec_anchor:
                prev_sec_anchor = doc.paragraphs[0] if doc.paragraphs else None

            for section in image_config:
                sec_id = section["id"]
                slots = section_to_slots_with_images.get(sec_id, [])
                anchor = section_anchors.get(sec_id)

                if not anchor:
                    if prev_sec_anchor:
                        p_title = doc.add_paragraph()
                        prev_sec_anchor._p.addnext(p_title._p)
                        r_title = p_title.add_run(section["title"])
                        r_title.bold = True
                        set_run_font(r_title, FONT_NAME_CN, 14)
                        anchor = p_title
                    else:
                        anchor = doc.paragraphs[-1] if doc.paragraphs else None

                for idx, slot_info in enumerate(slots):
                    if slot_info["id"] not in slot_anchors:
                        mark = chr(ord('a') + idx)
                        print(f"DEBUG: Inserting overflow slot {slot_info['id']} after anchor")
                        anchor = _add_image_slot_layout_v2(doc, anchor, slot_info["path"], mark, slot_info["label"], WORD_IMAGE_WIDTH_CM)

                if not sec_id.startswith("sec_") or len(sec_id) > 6:
                    analysis_text = analysis_results.get(sec_id, "none")
                    p_ana = doc.add_paragraph()
                    anchor._p.addnext(p_ana._p)
                    r_ana = p_ana.add_run(analysis_text)
                    set_run_font(r_ana, FONT_NAME_CN, 10.5)
                    anchor = p_ana

                prev_sec_anchor = anchor

        # -------------------------
        # 知识库检索（如果选择了知识库）
        # -------------------------
        kb_knowledge = ""
        if kb_id and KB_API_URL:
            yield sse("⏳ 正在检索知识库...")
            kb_token =kb_login()
            if kb_token:
                query_text =background_overview if background_overview and background_overview != "none" else project_name_val
                kb_knowledge=kb_retrieve(kb_token, kb_id, query_text, top_k=5, score_threshold=0.3)
                if kb_knowledge:
                    yield sse(f"⏳ 知识库检索完成，获取到 {len(kb_knowledge)} 字相关知识内容")
                else:
                    yield sse("⏳ 知识库未检索到相关内容，将不使用知识库增强")
            else:
                yield sse("⚠️ 知识库登录失败，将不使用知识库增强")

        # -------------------------
        # 结果生成与保存
        # -------------------------
        yield sse("⏳ 正在调用 LLM（生成结论/摘要）...")

        # 收集所有 VLM 分析结果，兼容不同 section ID 命名
        all_analyses = []
        for section in image_config:
            sec_id = section["id"]
            slots_with_img = section_to_slots_with_images.get(sec_id, [])
            if not slots_with_img:
                continue
            all_analyses.append({
                "title": section.get("title", ""),
                "result": analysis_results.get(sec_id, "none"),
            })

        final_prompt = build_final_llm_prompt(
            background_overview=background_overview,
            analyses=all_analyses,
            kb_knowledge=kb_knowledge,
        )
        try:
            llm_out = call_openai_chat(
                model=llm_model,
                messages=[{"role": "user", "content": final_prompt}],
                temperature=0.3,
                max_tokens=2000,
            )
            print(f"DEBUG: LLM raw output: {llm_out[:500]}")
            sections = parse_llm_json(llm_out)
            print(f"DEBUG: Parsed sections keys: {list(sections.keys())}")
            if not sections:
                print("WARNING: LLM JSON parsing returned empty result")
        except Exception as e:
            print(f"ERROR: LLM call failed: {e}\n{traceback.format_exc()}")
            sections = {}

        final_mapping = {
            PH_DISCUSSION: sections.get("分析与讨论", "") or "none",
            PH_CONCLUSION: sections.get("结论", "") or "none",
            PH_ABSTRACT: sections.get("报告摘要", "") or "none",
            PH_KEYWORDS: sections.get("关键词", "") or "none",
        }
        replace_text_placeholders_with_style(doc, final_mapping, FONT_NAME_CN, FONT_SIZE_5_PT)
        doc.save(str(report_path))

        # ── 计算页数并插入封面 ──
        yield sse("⏳ 正在计算页数...")
        page_count = _get_page_count_win32(str(report_path))
        if page_count:
            _insert_page_count(doc, page_count)
            doc.save(str(report_path))
            yield sse(f"  页数：{page_count}")

        msg = f"✅ 报告同步生成完成！\n- 章节总数: {len(image_config)}"
        yield sse(msg, file_path=str(report_path), done=True)

    except Exception as e:
        yield sse(f"❌ 运行失败：{str(e)}\n{traceback.format_exc()}", done=True)


# =========================
# Flask 应用
# =========================
app = Flask(__name__, static_folder='.', static_url_path='')


@app.route('/')
def index():
    resp = send_from_directory('.', 'index.html')
    resp.headers['Cache-Control'] = 'no-cache, no-store, must-revalidate'
    return resp


@app.route('/api/generate', methods=['POST'])
def api_generate():
    form = flask_request.form
    # 在主线程中先保存所有文件，因为线程开启后 request context 会失效
    files_map = {}
    for key, f in flask_request.files.items():
        if f and f.filename:
            tmp_save_path = _UPLOAD_TMP / f"{uuid.uuid4().hex}_{f.filename}"
            f.save(str(tmp_save_path))
            files_map[key] = tmp_save_path

    q: queue.Queue = queue.Queue()

    def worker():
        for chunk in _generate_report_stream(form, files_map):
            q.put(chunk)
        q.put(None)

    t = threading.Thread(target=worker, daemon=True)
    t.start()

    def stream():
        while True:
            item = q.get()
            if item is None:
                break
            yield item

    return Response(
        stream(),
        mimetype='text/event-stream',
        headers={'Cache-Control': 'no-cache', 'X-Accel-Buffering': 'no'},
    )


@app.route('/api/models', methods=['GET'])
def api_list_models():
    model_type = flask_request.args.get('model_type', '').strip().upper()

    url = f"{BASE_URL.rstrip('/')}/models"
    headers = {"Authorization": f"Bearer {API_KEY}", "Content-Type": "application/json"}

    try:
        r = requests.get(url, headers=headers, timeout=15)
        if r.status_code == 200:
            data = r.json()
            models = data.get("data", [])
            if model_type:
                models = [m for m in models if (m.get("type_code") or m.get("typeCode") or "").upper() == model_type]
            return {"data": models, "error": None}
        elif r.status_code in (401, 403):
            return {"error": "API Key 无效或无访问权限", "data": []}, r.status_code
        else:
            return {"error": f"服务异常 (HTTP {r.status_code})", "data": []}, 502
    except requests.ConnectionError:
        return {"error": "无法连接到服务器，请检查地址或网络", "data": []}, 502
    except requests.Timeout:
        return {"error": "请求超时，请稍后重试", "data": []}, 504
    except Exception as e:
        return {"error": f"请求失败: {str(e)}", "data": []}, 500


@app.route('/api/kb-list', methods=['GET'])
def api_kb_list():
    token = kb_login()
    if not token:
        return {"error": "知识库登录失败，请检查 KB 配置", "data": []}, 502
    kbs = kb_list(token)
    return {"data": kbs, "error": None}


@app.route('/api/download')
def api_download():
    file_path = flask_request.args.get('path', '')
    p = Path(file_path)
    if not p.exists() or not p.is_file():
        return 'File not found', 404
    try:
        p.resolve().relative_to(JOB_ROOT.resolve())
    except ValueError:
        return 'Forbidden', 403
    return send_file(str(p.resolve()), as_attachment=True, download_name=p.name)


if __name__ == '__main__':
    print("=" * 60)
    print("  失效报告生成助手")
    print("  访问地址：http://127.0.0.1:7860")
    print("=" * 60)
    app.run(host='0.0.0.0'
                 '', port=7860, debug=False, threaded=True)

