from __future__ import annotations

import base64
import io
import json
import os
import queue
import re
import shutil
import tempfile
import threading
import traceback
import uuid
from datetime import datetime
from pathlib import Path
from typing import Dict, Tuple, Any, List

import requests
from flask import Flask, request as flask_request, Response, send_file, send_from_directory
from PIL import Image, ImageOps
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


# =========================
# SpiritX OpenAI 兼容端点
# =========================
BASE_URL = "https://api.siliconflow.cn/v1"
VLM_MODEL = "Qwen/Qwen3-VL-32B-Instruct"
LLM_MODEL = "Qwen/Qwen3-32B"

API_KEY = "sk-daexzelyhxyxtbykxhgizoqdudxoiysexuxmguuqbzhnolal"

# =========================
# 输出目录
# =========================
JOB_ROOT = Path.home() / "Desktop" / "gradio_report_jobs"

# =========================
# 图片规则
# =========================
TARGET_W_PX = 782
TARGET_H_PX = 591
TARGET_SIZE = (TARGET_W_PX, TARGET_H_PX)

WORD_IMAGE_WIDTH_CM = 6.62  # 插入Word宽度（cm）

# =========================
# 字体要求：替换进去的文本统一宋体五号
# =========================
FONT_NAME_CN = "宋体"
FONT_SIZE_5_PT = 10.5  # 五号

# =========================
# 文本字段（用户填写）
# =========================
TEXT_FIELDS = [
    ("项目名称", "{{项目名称}}"),
    ("背景概述", "{{背景概述}}"),
]
OVERVIEW_SYNC_PLACEHOLDERS = ["{{背景概述}}", "{{概述}}"]

# =========================
# 图片槽位：slot_id -> (展示名, 分类, 固定尺寸?, 对应模板占位符)
# =========================
IMAGE_SLOTS: Dict[str, Tuple[str, str, bool, str]] = {
    "1.1": ("1.1 保险轴断裂部位", "1_断裂保险轴", True, "{{image_1.1}}"),
    "1.2": ("1.2 保险轴配套部件", "1_断裂保险轴", True, "{{image_1.2}}"),

    "2.1": ("2.1 限位凸起母体 15x", "2_宏观断口形貌", True, "{{image_2.1}}"),
    "2.2": ("2.2 限位凸起母体 20x", "2_宏观断口形貌", True, "{{image_2.2}}"),
    "2.3": ("2.3 限位凸起掉落部分 25x", "2_宏观断口形貌", True, "{{image_2.3}}"),
    "2.4": ("2.4 保险轴限位凸起掉落部分 25x", "2_宏观断口形貌", True, "{{image_2.4}}"),

    "3.1": ("3.1 宏观断口", "3_微观断口形貌", True, "{{image_3.1}}"),
    "3.2": ("3.2 裂纹源区", "3_微观断口形貌", True, "{{image_3.2}}"),
    "3.3": ("3.3 疲劳条带", "3_微观断口形貌", True, "{{image_3.3}}"),
    "3.4": ("3.4 扩展区 1", "3_微观断口形貌", True, "{{image_3.4}}"),
    "3.5": ("3.5 扩展区 2", "3_微观断口形貌", True, "{{image_3.5}}"),
    "3.6": ("3.6 瞬断区", "3_微观断口形貌", True, "{{image_3.6}}"),

    "4.1": ("4.1 显微组织 100x", "4_显微组织分析", True, "{{image_4.1}}"),
    "4.2": ("4.2 显微组织 500x", "4_显微组织分析", True, "{{image_4.2}}"),

    "5.1": ("5.1 化学分析表", "5_化学分析", False, "{{image_5.1}}"),
    "6.1": ("6.1 硬度试验表", "6_性能验证", False, "{{image_6.1}}"),
}

# =========================
# 模型生成替换目标占位符
# =========================
PH_BASIC = "{{基本图片分析}}"
PH_MACRO = "{{宏观断口图片分析}}"
PH_MICRO = "{{微观断口图片分析}}"
PH_MICROSTRUCT = "{{显微组织图片分析}}"
PH_CHEM = "{{化学分析}}"
PH_PERF = "{{性能验证}}"

PH_DISCUSSION = "{{分析与讨论}}"
PH_CONCLUSION = "{{结论}}"
PH_ABSTRACT = "{{报告摘要}}"
PH_KEYWORDS = "{{关键词}}"


# -------------------------
# 基础工具
# -------------------------
def ensure_dir(p: Path) -> Path:
    p.mkdir(parents=True, exist_ok=True)
    return p


def normalize_text(val: Any) -> str:
    if val is None:
        return "none"
    s = str(val).strip()
    return s if s else "none"


def safe_name(s: str, fallback: str = "project") -> str:
    s = (s or "").strip()
    if not s:
        return fallback
    s = re.sub(r"[^\w\u4e00-\u9fff\- ]+", "_", s)
    s = re.sub(r"\s+", "_", s).strip("_")
    return s or fallback


def find_template_docx() -> Path:
    local_desktop = Path.home() / "Desktop"
    icloud_desktop = Path.home() / "Library" / "Mobile Documents" / "com~apple~CloudDocs" / "Desktop"
    candidates = []
    for desk in [local_desktop, icloud_desktop]:
        for name in ["模版.docx", "模板.docx"]:
            candidates.append(desk / name)
    for p in candidates:
        if p.exists():
            return p
    raise FileNotFoundError("找不到模板文件，已尝试：\n" + "\n".join(str(x) for x in candidates))


def set_run_font(run, font_name: str, size_pt: float) -> None:
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    r = run._element.rPr
    if r is None:
        r = run._element.get_or_add_rPr()
    rFonts = r.rFonts
    if rFonts is None:
        rFonts = r.get_or_add_rFonts()
    rFonts.set(qn("w:eastAsia"), font_name)
    rFonts.set(qn("w:ascii"), font_name)
    rFonts.set(qn("w:hAnsi"), font_name)


# -------------------------
# 图片处理
# -------------------------
def to_rgb_image(src_path: Path) -> Image.Image:
    im = Image.open(src_path)
    im = ImageOps.exif_transpose(im)
    if im.mode == "RGBA":
        bg = Image.new("RGB", im.size, (255, 255, 255))
        bg.paste(im, mask=im.split()[-1])
        im = bg
    elif im.mode != "RGB":
        im = im.convert("RGB")
    return im


def aspect_fit_pad(im: Image.Image, target_size: Tuple[int, int]) -> Tuple[Image.Image, Dict[str, Any]]:
    tw, th = target_size
    ow, oh = im.size
    scale = min(tw / ow, th / oh)
    nw = max(1, int(round(ow * scale)))
    nh = max(1, int(round(oh * scale)))
    im_resized = im.resize((nw, nh), Image.Resampling.LANCZOS)
    canvas = Image.new("RGB", (tw, th), (255, 255, 255))
    x = (tw - nw) // 2
    y = (th - nh) // 2
    canvas.paste(im_resized, (x, y))
    info = {
        "original_px": {"w": ow, "h": oh},
        "resized_px": {"w": nw, "h": nh},
        "final_px": {"w": tw, "h": th},
        "method": "aspect-fit + white padding",
    }
    return canvas, info


def image_to_jpg_bytes(im: Image.Image, quality: int = 95) -> bytes:
    buf = io.BytesIO()
    im.save(buf, format="JPEG", quality=quality, optimize=True)
    return buf.getvalue()


def jpg_bytes_to_base64(data: bytes) -> str:
    return base64.b64encode(data).decode("utf-8")


def data_url_from_b64_jpg(b64: str) -> str:
    return "data:image/jpeg;base64," + b64


# -------------------------
# OpenAI 兼容调用
# -------------------------
def call_openai_chat(
    model: str,
    messages: List[Dict[str, Any]],
    temperature: float = 0.2,
    max_tokens: int = 800,
    timeout: int = 180,
) -> str:
    url = BASE_URL.rstrip("/") + "/chat/completions"
    headers = {"Authorization": f"Bearer {API_KEY}", "Content-Type": "application/json"}
    payload = {"model": model, "messages": messages, "temperature": temperature, "max_tokens": max_tokens}
    r = requests.post(url, headers=headers, json=payload, timeout=timeout)
    r.raise_for_status()
    data = r.json()
    return data["choices"][0]["message"]["content"]


# -------------------------
# Word 文本替换（替换进去的文字统一宋体五号）
# -------------------------
def _replace_in_paragraph_with_style(paragraph, mapping: Dict[str, str], font_name: str, size_pt: float) -> None:
    if not paragraph.runs:
        return
    full_text = "".join(run.text for run in paragraph.runs)
    if not any(k in full_text for k in mapping.keys()):
        return

    changed_any = False
    for run in paragraph.runs:
        original = run.text
        replaced = original
        for k, v in mapping.items():
            if k in replaced:
                replaced = replaced.replace(k, v)
        if replaced != original:
            run.text = replaced
            set_run_font(run, font_name, size_pt)
            changed_any = True

    new_full = "".join(run.text for run in paragraph.runs)
    if changed_any and not any(k in new_full for k in mapping.keys()):
        return

    replaced_all = full_text
    for k, v in mapping.items():
        replaced_all = replaced_all.replace(k, v)

    for run in paragraph.runs:
        run.text = ""

    run0 = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
    run0.text = replaced_all
    set_run_font(run0, font_name, size_pt)


def replace_text_placeholders_with_style(doc: Document, mapping: Dict[str, str], font_name: str, size_pt: float) -> None:
    for p in doc.paragraphs:
        _replace_in_paragraph_with_style(p, mapping, font_name, size_pt)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    _replace_in_paragraph_with_style(p, mapping, font_name, size_pt)


# -------------------------
# Word 图片替换
# -------------------------
def _insert_image_into_paragraph(paragraph, placeholder: str, image_path: Path, width_cm: float) -> bool:
    if placeholder not in paragraph.text:
        return False
    for run in paragraph.runs:
        run.text = ""
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=Cm(width_cm))
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return True


def replace_image_placeholders(doc: Document, placeholder_to_image: Dict[str, Path], width_cm: float) -> List[str]:
    remaining = set(placeholder_to_image.keys())
    for p in doc.paragraphs:
        for ph, img_path in list(placeholder_to_image.items()):
            if ph in p.text and _insert_image_into_paragraph(p, ph, img_path, width_cm):
                remaining.discard(ph)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for ph, img_path in list(placeholder_to_image.items()):
                        if ph in p.text and _insert_image_into_paragraph(p, ph, img_path, width_cm):
                            remaining.discard(ph)
    return sorted(list(remaining))


# -------------------------
# VLM Prompt：更短 + 允许标点 + 输出更精简（2~4 句）
# -------------------------
def build_vlm_prompt(background_overview: str, what: str) -> str:
    return (
        "你是一名材料失效分析工程师。\n"
        f"背景概述：{background_overview}\n"
        f"任务：{what}\n"
        "要求：只根据图片做客观描述与对比，不要编造，不要引用外部资料，不要下结论。\n"
        "输出：一段中文，2到4句话，允许使用标点，尽量简短。"
    )


def build_vlm_messages(prompt: str, image_b64_list: List[str]) -> List[Dict[str, Any]]:
    content = [{"type": "text", "text": prompt}]
    for b64 in image_b64_list:
        content.append({"type": "image_url", "image_url": {"url": data_url_from_b64_jpg(b64)}})
    return [{"role": "user", "content": content}]


# -------------------------
# LLM 最终生成：分析与讨论/结论/报告摘要/关键词（JSON）
# -------------------------
def build_final_llm_prompt(
    background_overview: str,
    basic: str,
    macro: str,
    micro: str,
    microstruct: str,
    chem: str,
    perf: str,
) -> str:
    return (
        "你是一名材料失效分析报告撰写专家。\n"
        "请基于背景概述与各部分图片分析内容，生成报告中的关键章节。\n"
        "必须输出严格的 JSON，不要包含多余文字。\n"
        "JSON keys 必须是：分析与讨论 结论 报告摘要 关键词\n"
        "要求：\n"
        "1 分析与讨论：1到2段，语言专业但不过长。\n"
        "2 结论：仅输出一整段中文（80到150字），必须包含标点符号，不要分条，不要使用 JSON 数组。\n"
        "3 报告摘要：一段话，80到150字。\n"
        "4 关键词：4到8个词，用空格分隔。\n"
        "\n背景概述：\n"
        f"{background_overview}\n"
        "\n基本图片分析：\n"
        f"{basic}\n"
        "\n宏观断口图片分析：\n"
        f"{macro}\n"
        "\n微观断口图片分析：\n"
        f"{micro}\n"
        "\n显微组织图片分析：\n"
        f"{microstruct}\n"
        "\n化学分析：\n"
        f"{chem}\n"
        "\n性能验证：\n"
        f"{perf}\n"
    )


def parse_llm_json(text: str) -> Dict[str, str]:
    def normalize_section(val: Any) -> str:
        if isinstance(val, list):
            items = [str(x).strip() for x in val if str(x).strip()]
            return "\n".join(f"{i}. {item}" for i, item in enumerate(items, start=1)) if items else ""
        if isinstance(val, dict):
            lines = []
            for k, v in val.items():
                ks = str(k).strip()
                vs = str(v).strip()
                if ks and vs:
                    lines.append(f"{ks}: {vs}")
            return "\n".join(lines)
        return str(val or "").strip()

    if not text:
        return {}
    s = text.strip()
    m = re.search(r"\{.*\}", s, flags=re.S)
    if m:
        s = m.group(0)
    try:
        obj = json.loads(s)
        return {
            "分析与讨论": normalize_section(obj.get("分析与讨论", "")),
            "结论": normalize_section(obj.get("结论", "")),
            "报告摘要": normalize_section(obj.get("报告摘要", "")),
            "关键词": normalize_section(obj.get("关键词", "")),
        }
    except Exception:
        return {}


# -------------------------
# 主流程（Flask SSE 流式版）
# -------------------------
_UPLOAD_TMP = Path(tempfile.gettempdir()) / "report_tool_uploads"
_UPLOAD_TMP.mkdir(parents=True, exist_ok=True)

LLM_CHOICES = [
    "Qwen/Qwen3-32B",
    "deepseek-ai/DeepSeek-R1-Distill-Qwen-32B",
    "Qwen/Qwen2.5-72B-Instruct",
]
VLM_CHOICES = [
    "Qwen/Qwen3-VL-32B-Instruct",
    "Qwen/Qwen2.5-VL-72B-Instruct",
    "Pro/Qwen/Qwen2.5-VL-7B-Instruct",
]
KB_CHOICES = [
    "工艺知识库",
    "材料知识库",
    "检测标准知识库",
]


def _generate_report_stream(form, files):
    """Generator: yields SSE data strings for Flask streaming response."""

    def sse(msg: str, file_path: str | None = None, done: bool = False) -> str:
        payload = {"msg": msg, "file": file_path, "done": done}
        return f"data: {json.dumps(payload, ensure_ascii=False)}\n\n"

    try:
        yield sse("⏳ 正在检查模板文件...")
        template_path = find_template_docx()

        project_name_val = form.get("项目名称", "")
        background_overview = form.get("背景概述", "") or "none"
        llm_model = form.get("llm_model", "") or LLM_MODEL
        vlm_model = form.get("vlm_model", "") or VLM_MODEL

        text_data = {
            "项目名称": normalize_text(project_name_val),
            "背景概述": normalize_text(background_overview),
        }

        ts = datetime.now().strftime("%Y%m%d_%H%M%S")
        proj = safe_name(text_data["项目名称"], fallback="project")
        job_id = f"{ts}_{proj}"
        job_dir = ensure_dir(JOB_ROOT / job_id)
        images_dir = ensure_dir(job_dir / "images_jpg")

        images_payload: Dict[str, Any] = {}
        placeholder_to_image_path: Dict[str, Path] = {}
        slot_to_b64: Dict[str, str] = {}
        total_images = 0
        total_bytes = 0

        yield sse("⏳ 正在处理上传图片...")
        total_slots = len(IMAGE_SLOTS)
        for idx, (slot_id, (display, category, fixed_size, placeholder)) in enumerate(
            IMAGE_SLOTS.items(), start=1
        ):
            uploaded = files.get(f"image_{slot_id}")
            if not uploaded or not uploaded.filename:
                images_payload[slot_id] = {
                    "display": display, "category": category, "provided": False,
                    "placeholder": placeholder, "jpg_relpath": None,
                    "jpg_base64": None, "process_info": None,
                }
                yield sse(f"⏳ 图片处理进度：{idx}/{total_slots}")
                continue

            tmp_path = _UPLOAD_TMP / f"{uuid.uuid4().hex}_{uploaded.filename}"
            uploaded.save(str(tmp_path))
            try:
                im = to_rgb_image(tmp_path)
                process_info = {
                    "uploaded_filename": uploaded.filename,
                    "rule": "1-4_fixed_size" if fixed_size else "5-6_original_size",
                    "original_px": {"w": im.size[0], "h": im.size[1]},
                }
                if fixed_size:
                    im, fit_info = aspect_fit_pad(im, TARGET_SIZE)
                    process_info.update(fit_info)
                else:
                    process_info.update({
                        "final_px": {"w": im.size[0], "h": im.size[1]},
                        "method": "convert-to-RGB-only",
                    })

                jpg_bytes = image_to_jpg_bytes(im, quality=95)
                b64 = jpg_bytes_to_base64(jpg_bytes)
                jpg_name = f"{slot_id}.jpg"
                jpg_path = images_dir / jpg_name
                jpg_path.write_bytes(jpg_bytes)
                total_images += 1
                total_bytes += len(jpg_bytes)
                images_payload[slot_id] = {
                    "display": display, "category": category, "provided": True,
                    "placeholder": placeholder,
                    "jpg_relpath": f"images_jpg/{jpg_name}",
                    "jpg_base64": b64, "process_info": process_info,
                }
                placeholder_to_image_path[placeholder] = jpg_path
                slot_to_b64[slot_id] = b64
            except Exception as e:
                images_payload[slot_id] = {
                    "display": display, "category": category, "provided": True,
                    "placeholder": placeholder, "jpg_relpath": None,
                    "jpg_base64": None, "error": str(e),
                }
            finally:
                try:
                    tmp_path.unlink()
                except Exception:
                    pass
            yield sse(f"⏳ 图片处理进度：{idx}/{total_slots}")

        payload_data = {
            "job_id": job_id,
            "created_at": datetime.now().isoformat(timespec="seconds"),
            "job_dir": str(job_dir),
            "text": text_data,
            "images": images_payload,
            "stats": {
                "images_saved": total_images,
                "approx_total_jpg_mb": round(total_bytes / 1024 / 1024, 2),
            },
        }
        payload_path = job_dir / "payload.json"
        payload_path.write_text(
            json.dumps(payload_data, ensure_ascii=False, indent=2), encoding="utf-8"
        )

        report_name = f"报告_{proj}_{ts}.docx"
        report_path = job_dir / report_name
        shutil.copyfile(template_path, report_path)
        doc = Document(str(report_path))

        yield sse("⏳ 正在替换模板文本与图片...")
        base_mapping = {placeholder: text_data[name] for name, placeholder in TEXT_FIELDS}
        for ph in OVERVIEW_SYNC_PLACEHOLDERS:
            base_mapping[ph] = background_overview
        replace_text_placeholders_with_style(doc, base_mapping, FONT_NAME_CN, FONT_SIZE_5_PT)
        remaining_images = replace_image_placeholders(
            doc, placeholder_to_image_path, width_cm=WORD_IMAGE_WIDTH_CM
        )

        if not API_KEY.strip():
            doc.save(str(report_path))
            warn = ""
            if remaining_images:
                warn = "\n⚠️ 以下图片占位符在模板中未找到：\n" + "\n".join(remaining_images)
            msg = (
                "✅ 已生成报告并完成文本与图片替换（未调用模型，因为 API_KEY 为空）\n"
                f"- 模板来源: {template_path}\n"
                f"- 报告文件: {report_path}\n"
                f"- payload.json: {payload_path}\n" + warn
            )
            yield sse(msg, file_path=str(report_path), done=True)
            return

        def run_vlm(slot_ids: List[str], task_desc: str) -> str:
            b64s = [slot_to_b64[sid] for sid in slot_ids if sid in slot_to_b64 and slot_to_b64[sid]]
            if not b64s:
                return "none"
            prompt = build_vlm_prompt(background_overview, task_desc)
            messages = build_vlm_messages(prompt, b64s)
            try:
                out = call_openai_chat(
                    model=vlm_model, messages=messages, temperature=0.2,
                    max_tokens=220, timeout=180,
                )
                return out.strip() if out and out.strip() else "none"
            except Exception:
                return "none"

        yield sse("⏳ 正在调用 VLM（图片分析）...")
        basic = run_vlm(["1.1", "1.2"], "观察保险轴断裂部位与配套部件，描述可见结构特征与异常点。")
        macro = run_vlm(["2.1", "2.2", "2.3", "2.4"], "观察宏观断口形貌，描述断口粗糙度、边缘特征、可能的裂纹起始迹象等。")
        micro = run_vlm(["3.1", "3.2", "3.3", "3.4", "3.5", "3.6"], "观察微观断口形貌，描述裂纹源区、疲劳条带、扩展区与瞬断区的形貌差异。")
        microstruct = run_vlm(["4.1", "4.2"], "观察显微组织，描述组织均匀性、晶粒特征、析出相或缺陷迹象等。")
        chem = run_vlm(["5.1"], "阅读化学分析表，概述主要元素组成特征与是否存在明显异常。")
        perf = run_vlm(["6.1"], "阅读硬度试验表，概述硬度分布是否均匀及是否存在明显偏离。")

        vlm_mapping = {
            PH_BASIC: basic, PH_MACRO: macro, PH_MICRO: micro,
            PH_MICROSTRUCT: microstruct, PH_CHEM: chem, PH_PERF: perf,
        }
        replace_text_placeholders_with_style(doc, vlm_mapping, FONT_NAME_CN, FONT_SIZE_5_PT)

        yield sse("⏳ 正在调用 LLM（生成结论/摘要）...")
        final_prompt = build_final_llm_prompt(
            background_overview=background_overview,
            basic=basic, macro=macro, micro=micro,
            microstruct=microstruct, chem=chem, perf=perf,
        )
        try:
            llm_out = call_openai_chat(
                model=llm_model,
                messages=[{"role": "user", "content": final_prompt}],
                temperature=0.35, max_tokens=1500, timeout=180,
            )
            sections = parse_llm_json(llm_out)
        except Exception:
            sections = {}

        final_mapping = {
            PH_DISCUSSION: sections.get("分析与讨论", "") or "none",
            PH_CONCLUSION: sections.get("结论", "") or "none",
            PH_ABSTRACT: sections.get("报告摘要", "") or "none",
            PH_KEYWORDS: sections.get("关键词", "") or "none",
        }
        replace_text_placeholders_with_style(doc, final_mapping, FONT_NAME_CN, FONT_SIZE_5_PT)
        doc.save(str(report_path))

        warn = ""
        if remaining_images:
            warn = "\n⚠️ 以下图片占位符在模板中未找到：\n" + "\n".join(remaining_images)
        msg = (
            "✅ 报告生成完成（已调用 VLM + LLM）\n"
            f"- 模板来源: {template_path}\n"
            f"- 报告文件: {report_path}\n"
            f"- payload.json: {payload_path}\n"
            f"- 已保存图片：{total_images} 张（images_jpg/）\n" + warn
        )
        yield sse(msg, file_path=str(report_path), done=True)

    except Exception as e:
        err = traceback.format_exc()
        yield sse(f"❌ 运行失败：{e}\n\n{err}", done=True)


# =========================
# Flask 应用
# =========================
app = Flask(__name__, static_folder='.', static_url_path='')


@app.route('/')
def index():
    return send_from_directory('.', 'index.html')


@app.route('/api/generate', methods=['POST'])
def api_generate():
    form = flask_request.form
    files = flask_request.files
    q: queue.Queue = queue.Queue()

    def worker():
        for chunk in _generate_report_stream(form, files):
            q.put(chunk)
        q.put(None)

    t = threading.Thread(target=worker, daemon=True)
    t.start()

    def stream():
        while True:
            item = q.get()
            if item is None:
                break
            yield item

    return Response(
        stream(),
        mimetype='text/event-stream',
        headers={'Cache-Control': 'no-cache', 'X-Accel-Buffering': 'no'},
    )


@app.route('/api/download')
def api_download():
    file_path = flask_request.args.get('path', '')
    p = Path(file_path)
    if not p.exists() or not p.is_file():
        return 'File not found', 404
    try:
        p.resolve().relative_to(JOB_ROOT.resolve())
    except ValueError:
        return 'Forbidden', 403
    return send_file(str(p.resolve()), as_attachment=True, download_name=p.name)


if __name__ == '__main__':
    print("=" * 60)
    print("  失效报告生成助手")
    print("  访问地址：http://127.0.0.1:7860")
    print("=" * 60)
    app.run(host='127.0.0.1', port=7860, debug=False, threaded=True)

